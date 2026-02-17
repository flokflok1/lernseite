"""
LernsystemX File Context Service

Core text extraction from files for AI context:
- PDF extraction (via PDFService)
- DOCX extraction (python-docx)
- TXT/Markdown support

Higher-level operations (multi-file context, AI formatting,
preview, validation) are in context_part2.py.

Phase D4 - Universal KI-Authoring-System
"""

import io
import logging
from typing import Dict, Any
from pathlib import Path

# DOCX Library
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.application.services.pdf import PDFService, PDFExtractionError

logger = logging.getLogger(__name__)


class FileContextError(Exception):
    """Base exception for file context errors"""
    pass


class UnsupportedFileTypeError(FileContextError):
    """Raised for unsupported file types"""
    pass


class FileContextService:
    """
    File Context Service for AI Authoring.

    Extracts and combines text from multiple files to provide
    context for AI content generation.

    Supported formats:
    - PDF (via PDFService)
    - DOCX (Microsoft Word)
    - TXT (plain text)
    - MD (Markdown)

    Usage:
        >>> context = FileContextService.get_file_context(['file_id_1', 'file_id_2'])
        >>> print(context['combined_text'])
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    MAX_CONTEXT_LENGTH = 100000  # Characters

    @classmethod
    def extract_text(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from a file based on its type.

        Args:
            file_content: Raw file bytes
            filename: Original filename (used to detect type)

        Returns:
            Dict with:
            - text: Extracted text
            - filename: str
            - file_type: str (pdf, docx, txt, md)
            - word_count: int
            - char_count: int
            - metadata: dict (file-specific metadata)

        Raises:
            UnsupportedFileTypeError: For unsupported file types
            FileContextError: For extraction errors
        """
        ext = Path(filename).suffix.lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        try:
            if ext == '.pdf':
                return cls._extract_from_pdf(file_content, filename)
            elif ext == '.docx':
                return cls._extract_from_docx(file_content, filename)
            elif ext in {'.txt', '.md', '.markdown'}:
                return cls._extract_from_text(file_content, filename, ext)
            else:
                raise UnsupportedFileTypeError(f"No handler for: {ext}")
        except (PDFExtractionError, UnsupportedFileTypeError):
            raise
        except Exception as e:
            logger.error(f"File extraction error for {filename}: {str(e)}")
            raise FileContextError(f"Failed to extract text: {str(e)}")

    @classmethod
    def _extract_from_pdf(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF using PDFService."""
        result = PDFService.extract_text(file_content, filename)

        text = result.get('extracted_text', '')

        return {
            'text': text,
            'filename': filename,
            'file_type': 'pdf',
            'word_count': len(text.split()),
            'char_count': len(text),
            'metadata': {
                'page_count': result.get('page_count', 0),
                'pdf_metadata': result.get('extracted_metadata', {}),
                'structure': result.get('structure_analysis', {})
            }
        }

    @classmethod
    def _extract_from_docx(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise FileContextError("python-docx is not installed")

        try:
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)
        except Exception as e:
            raise FileContextError(f"Cannot read DOCX: {filename} - {str(e)}")

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    tables_text.append(' | '.join(row_text))

        # Combine all text
        full_text = '\n\n'.join(paragraphs)
        if tables_text:
            full_text += '\n\n--- Tabellen ---\n\n' + '\n'.join(tables_text)

        # Extract metadata
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else ''
            }
            metadata = {k: v for k, v in metadata.items() if v}

        return {
            'text': full_text,
            'filename': filename,
            'file_type': 'docx',
            'word_count': len(full_text.split()),
            'char_count': len(full_text),
            'metadata': {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'docx_properties': metadata
            }
        }

    @classmethod
    def _extract_from_text(cls, file_content: bytes, filename: str, ext: str) -> Dict[str, Any]:
        """Extract text from TXT/Markdown file."""
        # Try different encodings
        text = None
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise FileContextError(f"Cannot decode text file: {filename}")

        file_type = 'markdown' if ext in {'.md', '.markdown'} else 'txt'

        return {
            'text': text,
            'filename': filename,
            'file_type': file_type,
            'word_count': len(text.split()),
            'char_count': len(text),
            'metadata': {
                'line_count': text.count('\n') + 1
            }
        }
