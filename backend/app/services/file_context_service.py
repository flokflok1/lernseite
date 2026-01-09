"""
LernsystemX File Context Service

Extracts text content from files for AI context:
- PDF extraction (via PDFService)
- DOCX extraction (python-docx)
- TXT/Markdown support
- Combined multi-file context

Phase D4 - Universal KI-Authoring-System
"""

import io
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path

# DOCX Library
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.services.pdf import PDFService, PDFExtractionError
from app.repositories.courses.files import CourseFileRepository

logger = logging.getLogger(__name__)


class FileContextError(Exception):
    """Base exception for file context errors"""
    pass


class UnsupportedFileTypeError(FileContextError):
    """Raised for unsupported file types"""
    pass


class FileContextService:
    """
    File Context Service for AI Authoring.

    Extracts and combines text from multiple files to provide
    context for AI content generation.

    Supported formats:
    - PDF (via PDFService)
    - DOCX (Microsoft Word)
    - TXT (plain text)
    - MD (Markdown)

    Usage:
        >>> context = FileContextService.get_file_context(['file_id_1', 'file_id_2'])
        >>> print(context['combined_text'])
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.markdown'}
    MAX_CONTEXT_LENGTH = 100000  # Characters

    @classmethod
    def extract_text(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from a file based on its type.

        Args:
            file_content: Raw file bytes
            filename: Original filename (used to detect type)

        Returns:
            Dict with:
            - text: Extracted text
            - filename: str
            - file_type: str (pdf, docx, txt, md)
            - word_count: int
            - char_count: int
            - metadata: dict (file-specific metadata)

        Raises:
            UnsupportedFileTypeError: For unsupported file types
            FileContextError: For extraction errors
        """
        ext = Path(filename).suffix.lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        try:
            if ext == '.pdf':
                return cls._extract_from_pdf(file_content, filename)
            elif ext == '.docx':
                return cls._extract_from_docx(file_content, filename)
            elif ext in {'.txt', '.md', '.markdown'}:
                return cls._extract_from_text(file_content, filename, ext)
            else:
                raise UnsupportedFileTypeError(f"No handler for: {ext}")
        except (PDFExtractionError, UnsupportedFileTypeError):
            raise
        except Exception as e:
            logger.error(f"File extraction error for {filename}: {str(e)}")
            raise FileContextError(f"Failed to extract text: {str(e)}")

    @classmethod
    def _extract_from_pdf(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF using PDFService."""
        result = PDFService.extract_text(file_content, filename)

        text = result.get('extracted_text', '')

        return {
            'text': text,
            'filename': filename,
            'file_type': 'pdf',
            'word_count': len(text.split()),
            'char_count': len(text),
            'metadata': {
                'page_count': result.get('page_count', 0),
                'pdf_metadata': result.get('extracted_metadata', {}),
                'structure': result.get('structure_analysis', {})
            }
        }

    @classmethod
    def _extract_from_docx(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise FileContextError("python-docx is not installed")

        try:
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)
        except Exception as e:
            raise FileContextError(f"Cannot read DOCX: {filename} - {str(e)}")

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    tables_text.append(' | '.join(row_text))

        # Combine all text
        full_text = '\n\n'.join(paragraphs)
        if tables_text:
            full_text += '\n\n--- Tabellen ---\n\n' + '\n'.join(tables_text)

        # Extract metadata
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else ''
            }
            metadata = {k: v for k, v in metadata.items() if v}

        return {
            'text': full_text,
            'filename': filename,
            'file_type': 'docx',
            'word_count': len(full_text.split()),
            'char_count': len(full_text),
            'metadata': {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'docx_properties': metadata
            }
        }

    @classmethod
    def _extract_from_text(cls, file_content: bytes, filename: str, ext: str) -> Dict[str, Any]:
        """Extract text from TXT/Markdown file."""
        # Try different encodings
        text = None
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise FileContextError(f"Cannot decode text file: {filename}")

        file_type = 'markdown' if ext in {'.md', '.markdown'} else 'txt'

        return {
            'text': text,
            'filename': filename,
            'file_type': file_type,
            'word_count': len(text.split()),
            'char_count': len(text),
            'metadata': {
                'line_count': text.count('\n') + 1
            }
        }

    @classmethod
    def get_file_context(
        cls,
        file_ids: List[str],
        max_length: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Get combined text context from multiple files.

        Args:
            file_ids: List of file IDs from course_files table
            max_length: Optional max total character length

        Returns:
            Dict with:
            - combined_text: str (all files combined)
            - files: List of file info dicts
            - total_word_count: int
            - total_char_count: int
            - truncated: bool (if max_length exceeded)
        """
        if not file_ids:
            return {
                'combined_text': '',
                'files': [],
                'total_word_count': 0,
                'total_char_count': 0,
                'truncated': False
            }

        max_len = max_length or cls.MAX_CONTEXT_LENGTH
        files_info = []
        combined_parts = []
        total_chars = 0
        truncated = False

        for file_id in file_ids:
            try:
                # Get file from database
                file_record = CourseFileRepository.find_by_id(file_id)
                if not file_record:
                    logger.warning(f"File not found: {file_id}")
                    continue

                # Get file content
                file_path = file_record.get('file_path')
                filename = file_record.get('original_filename', 'unknown')

                if not file_path:
                    logger.warning(f"No file path for: {file_id}")
                    continue

                # Read file content
                try:
                    with open(file_path, 'rb') as f:
                        file_content = f.read()
                except FileNotFoundError:
                    logger.warning(f"File not found on disk: {file_path}")
                    continue

                # Extract text
                extraction = cls.extract_text(file_content, filename)
                text = extraction['text']

                # Check length limits
                if total_chars + len(text) > max_len:
                    remaining = max_len - total_chars
                    if remaining > 500:  # Only add if meaningful amount left
                        text = text[:remaining] + '\n\n[... Text abgeschnitten ...]'
                        extraction['char_count'] = remaining
                        truncated = True
                    else:
                        truncated = True
                        break

                total_chars += len(text)

                # Add file header
                header = f"\n\n=== Datei: {filename} ===\n\n"
                combined_parts.append(header + text)

                files_info.append({
                    'file_id': file_id,
                    'filename': filename,
                    'file_type': extraction['file_type'],
                    'word_count': extraction['word_count'],
                    'char_count': extraction['char_count'],
                    'metadata': extraction.get('metadata', {})
                })

            except Exception as e:
                logger.error(f"Error processing file {file_id}: {str(e)}")
                files_info.append({
                    'file_id': file_id,
                    'filename': 'unknown',
                    'error': str(e)
                })

        combined_text = ''.join(combined_parts).strip()

        return {
            'combined_text': combined_text,
            'files': files_info,
            'total_word_count': len(combined_text.split()),
            'total_char_count': len(combined_text),
            'truncated': truncated
        }

    @classmethod
    def extract_for_ai_context(
        cls,
        file_ids: List[str],
        context_type: str = 'general'
    ) -> str:
        """
        Extract and format file content for AI prompts.

        Args:
            file_ids: List of file IDs
            context_type: Type of context ('general', 'chapter', 'lesson', 'task')

        Returns:
            Formatted context string for AI prompts
        """
        context_data = cls.get_file_context(file_ids)

        if not context_data['combined_text']:
            return ""

        # Format header based on context type
        headers = {
            'general': "Referenz-Material",
            'chapter': "Kapitel-Quelldokumente",
            'lesson': "Lektions-Referenzmaterial",
            'task': "Aufgaben-Kontext"
        }

        header = headers.get(context_type, "Referenz-Material")

        # Build context string
        parts = [f"<{header.lower().replace(' ', '_').replace('-', '_')}>"]

        # Add file summary
        if context_data['files']:
            parts.append(f"\nEnthaltene Dateien ({len(context_data['files'])}):")
            for f in context_data['files']:
                if 'error' not in f:
                    parts.append(f"- {f['filename']} ({f['file_type']}, {f['word_count']} Wörter)")

        parts.append("\n--- Inhalt ---\n")
        parts.append(context_data['combined_text'])

        if context_data['truncated']:
            parts.append("\n\n[Hinweis: Text wurde aufgrund der Länge gekürzt]")

        parts.append(f"\n</{header.lower().replace(' ', '_').replace('-', '_')}>")

        return '\n'.join(parts)

    @classmethod
    def get_file_preview(
        cls,
        file_id: str,
        max_chars: int = 2000
    ) -> Dict[str, Any]:
        """
        Get a quick preview of a file.

        Args:
            file_id: File ID
            max_chars: Maximum characters for preview

        Returns:
            Dict with preview text and basic info
        """
        try:
            file_record = CourseFileRepository.find_by_id(file_id)
            if not file_record:
                return {'error': 'File not found', 'preview': ''}

            file_path = file_record.get('file_path')
            filename = file_record.get('original_filename', 'unknown')

            if not file_path:
                return {'error': 'No file path', 'preview': ''}

            with open(file_path, 'rb') as f:
                file_content = f.read()

            extraction = cls.extract_text(file_content, filename)
            text = extraction['text']

            preview = text[:max_chars]
            if len(text) > max_chars:
                preview += '\n\n[...]'

            return {
                'preview': preview,
                'filename': filename,
                'file_type': extraction['file_type'],
                'word_count': extraction['word_count'],
                'char_count': extraction['char_count'],
                'truncated': len(text) > max_chars
            }

        except Exception as e:
            logger.error(f"Preview error for {file_id}: {str(e)}")
            return {'error': str(e), 'preview': ''}

    @classmethod
    def validate_file(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate if a file can be processed.

        Args:
            file_content: Raw file bytes
            filename: Original filename

        Returns:
            Dict with:
            - valid: bool
            - file_type: str
            - error: Optional[str]
        """
        ext = Path(filename).suffix.lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            return {
                'valid': False,
                'file_type': ext,
                'error': f"Unsupported file type. Supported: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            }

        # Type-specific validation
        try:
            if ext == '.pdf':
                is_valid, error = PDFService.validate_pdf(file_content)
                return {
                    'valid': is_valid,
                    'file_type': 'pdf',
                    'error': error
                }
            elif ext == '.docx':
                # Try to open as DOCX
                if not DOCX_AVAILABLE:
                    return {
                        'valid': False,
                        'file_type': 'docx',
                        'error': 'DOCX support not available'
                    }
                try:
                    doc_file = io.BytesIO(file_content)
                    DocxDocument(doc_file)
                    return {'valid': True, 'file_type': 'docx', 'error': None}
                except Exception as e:
                    return {
                        'valid': False,
                        'file_type': 'docx',
                        'error': f"Invalid DOCX file: {str(e)}"
                    }
            else:
                # Text files - check if decodable
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        file_content.decode(encoding)
                        return {
                            'valid': True,
                            'file_type': 'txt' if ext == '.txt' else 'markdown',
                            'error': None
                        }
                    except UnicodeDecodeError:
                        continue
                return {
                    'valid': False,
                    'file_type': 'txt',
                    'error': 'Cannot decode text file'
                }

        except Exception as e:
            return {
                'valid': False,
                'file_type': ext,
                'error': str(e)
            }
